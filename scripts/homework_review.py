"""Build review documents from Chaoxing homework exports."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import time
import urllib.error
import urllib.request
from pathlib import Path
from typing import Callable, Iterable
from xml.sax.saxutils import escape


DEFAULT_BASE_URL = "https://api.deepseek.com"
DEFAULT_MODEL = "deepseek-v4-flash"
DEFAULT_MAX_TOKENS = 2000
DEFAULT_DOCX_FONT = "Microsoft YaHei"
QUESTION_FONT_SIZE_PT = 13
CALLOUT_FONT_SIZE_PT = 12
TIP_SHADE = "FFF4D6"
LINE_COLOR = "000000"


def normalize_question_text(text: str) -> str:
    text = re.sub(r"^\s*\d+[.．、]\s*", "", text or "")
    return re.sub(r"\s+", " ", text).strip()


def normalize_option(option: str) -> str:
    option = re.sub(r"\s+", " ", option or "").strip()
    option = re.sub(r"^([A-Z])\s*[.．、]\s*", r"\1. ", option)
    option = re.sub(r"^([A-Z])\. \1\s*[.．、]\s*", r"\1. ", option)
    return option


def normalize_answer(answer: object) -> str:
    text = str(answer or "").strip()
    parts = [part.strip() for part in text.split("###") if part.strip()]
    return "；".join(parts) if parts else text


def normalize_question(raw: dict, meta: dict | None = None, source_file: str = "") -> dict:
    meta = meta or {}
    question = dict(raw)
    question["courseName"] = question.get("courseName") or meta.get("courseName", "")
    if source_file:
        question["sourceFile"] = source_file
    question["question"] = normalize_question_text(question.get("question", ""))
    question["options"] = [
        normalize_option(option) for option in question.get("options", []) if option
    ]
    question["answer"] = normalize_answer(question.get("answer", ""))
    return question


def question_key(question: dict) -> str:
    payload = {
        "type": question.get("type", ""),
        "question": normalize_question_text(question.get("question", "")),
        "options": [normalize_option(item) for item in question.get("options", [])],
        "answer": question.get("answer", ""),
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def load_questions(input_path: Path | str) -> list[dict]:
    root = Path(input_path)
    files = sorted(root.rglob("*.json")) if root.is_dir() else [root]
    files = [file_path for file_path in files if is_source_json_file(file_path)]
    questions: list[dict] = []
    seen: set[str] = set()

    for file_path in files:
        if file_path.name.endswith(".enriched.json"):
            continue
        data = json.loads(file_path.read_text(encoding="utf-8-sig"))
        meta = data.get("meta", {}) if isinstance(data, dict) else {}
        raw_questions = data.get("questions", data if isinstance(data, list) else [])
        for raw in raw_questions:
            question = normalize_question(raw, meta, file_path.name)
            key = question_key(question)
            if key in seen:
                continue
            question["id"] = key[:12]
            seen.add(key)
            questions.append(question)

    return questions


def is_source_json_file(file_path: Path) -> bool:
    generated_names = {
        "questions.enriched.json",
        "questions.partial.json",
        "explanations.cache.json",
    }
    generated_dirs = {
        "output",
        "final-output",
        "progress-test-output",
        "layout-test-output",
        "verify-test-output",
        "review-card-test-output",
        "student-review-test-output",
        "json-mode-test-output",
        "api-test-output",
    }
    if file_path.name in generated_names:
        return False
    return not any(part in generated_dirs for part in file_path.parts)


def apply_limit(questions: list[dict], limit: int | None) -> list[dict]:
    if not limit or limit <= 0:
        return questions
    return questions[:limit]


def log_progress(
    stage: str,
    index: int,
    total: int,
    question: str,
    logger: Callable[[str], None] = print,
) -> None:
    title = re.sub(r"\s+", " ", question or "").strip()
    if len(title) > 48:
        title = title[:45] + "..."
    logger(f"[{index}/{total}] {stage}：{title}", flush=True) if logger is print else logger(
        f"[{index}/{total}] {stage}：{title}"
    )


def load_cache(cache_path: Path | None) -> dict[str, dict]:
    if not cache_path or not cache_path.exists():
        return {}
    return json.loads(cache_path.read_text(encoding="utf-8"))


def save_json(path: Path, data: object) -> None:
    save_json_atomic(path, data)


def save_json_atomic(path: Path, data: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = path.with_suffix(path.suffix + ".tmp")
    temp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    temp_path.replace(path)


def docx_font_name() -> str:
    font_name = os.getenv("DOCX_FONT", "").strip()
    if is_valid_font_name(font_name):
        return font_name
    return DEFAULT_DOCX_FONT


def is_valid_font_name(font_name: str) -> bool:
    if not font_name or len(font_name) > 80:
        return False
    return bool(re.fullmatch(r"[\w\s\-\u4e00-\u9fff]+", font_name, re.UNICODE))


def load_dotenv(path: Path | str = ".env") -> None:
    env_path = Path(path)
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8-sig").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def build_prompt(question: dict) -> list[dict]:
    options = "\n".join(question.get("options", [])) or "无"
    user = f"""题型：{question.get("type", "")}
题目：{question.get("question", "")}
选项：
{options}
已知正确答案：{question.get("answer", "")}

请生成适合复习背诵的中文解析，并严格输出 json。"""
    return [
        {
            "role": "system",
            "content": (
                "你是课程复习助手。根据题目、选项和已知正确答案生成解析，"
                "不要更改答案。必须只输出合法 json，不要输出 Markdown。"
                "所有内容都要面向正在复习的学生，像老师讲解这道题一样。"
                "json 格式示例："
                '{"correct_reason":"为什么正确答案正确",'
                '"wrong_options":[{"option":"A","reason":"为什么不选，干扰点在哪"}],'
                '"review_tip":"复习抓手：看到哪些关键词或条件时应如何快速判断",'
                '"knowledge_points":["这道题涉及的知识点，需展开说明其含义和复习时要抓住什么"],'
                '"principles":["这道题背后的判断原理，需说明遇到同类题怎么判断"]}。'
                "correct_reason 控制在 80 到 150 字；wrong_options 覆盖明显错误选项；"
                "review_tip 用 1 句话说明考试时如何快速识别答案；"
                "knowledge_points 和 principles 各给 1 到 4 条，每条 50 到 110 字；"
                "不要只写标签或名词短语，要展开说明为什么这个知识点与题目有关，"
                "以及学生复习时应该记住的判断线索。"
            ),
        },
        {"role": "user", "content": user},
    ]


def api_settings_from_env() -> dict[str, str]:
    api_key = (
        os.getenv("AI_API_KEY")
        or os.getenv("DEEPSEEK_API_KEY")
        or os.getenv("OPENAI_API_KEY")
    )
    if not api_key:
        raise RuntimeError("Missing AI_API_KEY, DEEPSEEK_API_KEY, or OPENAI_API_KEY.")
    return {
        "api_key": api_key,
        "base_url": os.getenv("AI_BASE_URL")
        or os.getenv("DEEPSEEK_BASE_URL")
        or DEFAULT_BASE_URL,
        "model": os.getenv("AI_MODEL") or os.getenv("DEEPSEEK_MODEL") or DEFAULT_MODEL,
    }


def call_chat_completion(messages: list[dict]) -> str:
    settings = api_settings_from_env()
    url = settings["base_url"].rstrip("/") + "/chat/completions"
    body: dict = {
        "model": settings["model"],
        "messages": messages,
        "stream": False,
        "max_tokens": int(os.getenv("AI_MAX_TOKENS", str(DEFAULT_MAX_TOKENS))),
        "temperature": float(os.getenv("AI_TEMPERATURE", "0.2")),
        "response_format": {"type": "json_object"},
    }
    thinking = os.getenv("AI_THINKING", "disabled")
    if "deepseek.com" in settings["base_url"] or os.getenv("AI_THINKING"):
        body["thinking"] = {"type": thinking}

    payload = json.dumps(body, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Authorization": f"Bearer {settings['api_key']}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    for attempt in range(3):
        try:
            with urllib.request.urlopen(request, timeout=60) as response:
                data = json.loads(response.read().decode("utf-8"))
            content = data["choices"][0]["message"]["content"].strip()
            if not content:
                raise RuntimeError("API returned empty JSON content.")
            return content
        except urllib.error.HTTPError as exc:
            details = exc.read().decode("utf-8", errors="replace")
            if exc.code < 500 or attempt == 2:
                raise RuntimeError(f"API request failed: HTTP {exc.code} {details}") from exc
        except urllib.error.URLError as exc:
            if attempt == 2:
                raise RuntimeError(f"API request failed: {exc}") from exc
        time.sleep(2**attempt)

    raise RuntimeError("API request failed after retries.")


def parse_explanation_response(content: str) -> dict:
    data = json.loads(content)
    return normalize_explanation(data)


def build_answer_check_prompt(question: dict) -> list[dict]:
    options = "\n".join(question.get("options", [])) or "无"
    user = f"""题型：{question.get("type", "")}
题目：{question.get("question", "")}
选项：
{options}
导出答案：{question.get("answer", "")}

请独立校验导出答案是否合理，并严格输出 json。"""
    return [
        {
            "role": "system",
            "content": (
                "你是答案校验器。请先独立判断题目的合理答案，再和导出答案比较。"
                "不要默认相信导出答案，也不要自动修改答案。必须只输出合法 json，"
                "不要输出 Markdown。json 格式示例："
                '{"provided_answer":"导出答案","model_answer":"你独立判断的答案",'
                '"verdict":"agree|disagree|uncertain",'
                '"confidence":0.0,'
                '"risk_level":"low|medium|high",'
                '"reason":"简要说明为什么一致、不一致或不确定",'
                '"needs_review":false}。'
                "如果题目有歧义、多个选项可能成立、信息不足或你不确定，"
                "verdict 用 uncertain，needs_review 用 true。"
            ),
        },
        {"role": "user", "content": user},
    ]


def parse_answer_check_response(content: str) -> dict:
    return normalize_answer_check(json.loads(content))


def normalize_answer_check(value: object) -> dict:
    if not isinstance(value, dict):
        value = {}
    verdict = str(value.get("verdict", "uncertain")).strip().lower()
    if verdict not in {"agree", "disagree", "uncertain"}:
        verdict = "uncertain"
    risk_level = str(value.get("risk_level", "")).strip().lower()
    if risk_level not in {"low", "medium", "high"}:
        risk_level = {"agree": "low", "disagree": "high"}.get(verdict, "medium")
    try:
        confidence = float(value.get("confidence", 0))
    except (TypeError, ValueError):
        confidence = 0.0
    needs_review = bool(value.get("needs_review", verdict != "agree" or risk_level != "low"))
    return {
        "provided_answer": str(value.get("provided_answer", "")).strip(),
        "model_answer": str(value.get("model_answer", "")).strip(),
        "verdict": verdict,
        "confidence": max(0.0, min(1.0, confidence)),
        "risk_level": risk_level,
        "reason": str(value.get("reason", "")).strip(),
        "needs_review": needs_review,
    }


def normalize_explanation(value: object) -> dict:
    if isinstance(value, dict):
        wrong_options = value.get("wrong_options", [])
        if not isinstance(wrong_options, list):
            wrong_options = []
        normalized_wrong_options = []
        for item in wrong_options:
            if isinstance(item, dict):
                normalized_wrong_options.append(
                    {
                        "option": str(item.get("option", "")).strip(),
                        "reason": str(item.get("reason", "")).strip(),
                    }
                )
            elif item:
                normalized_wrong_options.append({"option": "", "reason": str(item).strip()})

        return {
            "correct_reason": str(value.get("correct_reason", "")).strip(),
            "wrong_options": [
                item for item in normalized_wrong_options if item["option"] or item["reason"]
            ],
            "review_tip": str(value.get("review_tip", "")).strip(),
            "knowledge_points": _string_list(value.get("knowledge_points", [])),
            "principles": _string_list(value.get("principles", [])),
        }
    return {
        "correct_reason": str(value or "").strip(),
        "wrong_options": [],
        "review_tip": "",
        "knowledge_points": [],
        "principles": [],
    }


def _string_list(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if value:
        return [str(value).strip()]
    return []


def enrich_questions(
    questions: Iterable[dict],
    client: Callable[[list[dict]], str] = call_chat_completion,
    cache: dict[str, dict] | None = None,
    dry_run: bool = False,
    verify_answers: bool = False,
    cache_writer: Callable[[dict[str, dict], list[dict]], None] | None = None,
    logger: Callable[[str], None] = print,
) -> list[dict]:
    cache = cache or {}
    enriched: list[dict] = []
    question_list = list(questions)
    total = len(question_list)
    for index, question in enumerate(question_list, 1):
        item = dict(question)
        key = question_key(item)
        cached = cache.get(key)
        if cached and cached.get("explanation"):
            log_progress("使用缓存", index, total, item.get("question", ""), logger)
            item["explanation"] = normalize_explanation(cached["explanation"])
            item["cached_explanation_source"] = cached.get("explanation_source", "unknown")
            item["explanation_source"] = "cache"
            if cached.get("answer_check"):
                item["answer_check"] = normalize_answer_check(cached["answer_check"])
        elif item.get("explanation"):
            log_progress("使用原解析", index, total, item.get("question", ""), logger)
            item["explanation"] = normalize_explanation(item["explanation"])
            item["explanation_source"] = item.get("explanation_source", "platform")
        elif dry_run:
            log_progress("dry-run 占位", index, total, item.get("question", ""), logger)
            item["explanation"] = normalize_explanation(
                {"correct_reason": "待生成：dry-run 模式未调用 AI API。"}
            )
            item["explanation_source"] = "missing"
        else:
            log_progress("生成解析", index, total, item.get("question", ""), logger)
            try:
                item["explanation"] = parse_explanation_response(client(build_prompt(item)))
                item["explanation_source"] = "ai"
            except Exception as exc:
                item["explanation"] = normalize_explanation(
                    {
                        "correct_reason": (
                            "生成解析失败。请查看 processing_error 字段后重试，"
                            "或人工补充本题解析。"
                        )
                    }
                )
                item["explanation_source"] = "failed"
                item["processing_error"] = str(exc)
                log_progress("解析失败", index, total, item.get("question", ""), logger)
        if verify_answers and not item.get("answer_check") and item.get("explanation_source") != "failed":
            log_progress("答案校验", index, total, item.get("question", ""), logger)
            try:
                item["answer_check"] = parse_answer_check_response(
                    client(build_answer_check_prompt(item))
                )
            except Exception as exc:
                item["answer_check"] = normalize_answer_check(
                    {
                        "provided_answer": item.get("answer", ""),
                        "model_answer": "",
                        "verdict": "uncertain",
                        "confidence": 0,
                        "risk_level": "medium",
                        "reason": f"答案校验失败：{exc}",
                        "needs_review": True,
                    }
                )
                item["processing_error"] = (
                    f"{item.get('processing_error', '')}\n答案校验失败：{exc}"
                ).strip()
                log_progress("校验失败", index, total, item.get("question", ""), logger)
        enriched.append(item)
        if cache_writer:
            cache = update_cache(cache, [item])
            cache_writer(cache, enriched)
    return enriched


def update_cache(cache: dict[str, dict], questions: Iterable[dict]) -> dict[str, dict]:
    updated = dict(cache)
    for question in questions:
        if question.get("explanation") and question.get("explanation_source") != "missing":
            cached_question = {
                "explanation": normalize_explanation(question["explanation"]),
                "explanation_source": question.get("explanation_source", "ai"),
            }
            if question.get("answer_check"):
                cached_question["answer_check"] = normalize_answer_check(question["answer_check"])
            updated[question_key(question)] = cached_question
    return updated


def build_run_summary(questions: list[dict]) -> dict[str, int]:
    summary = {
        "total": len(questions),
        "cache": 0,
        "ai": 0,
        "platform": 0,
        "missing": 0,
        "failed": 0,
        "review_needed": 0,
    }
    for question in questions:
        source = question.get("explanation_source", "missing")
        if source in summary:
            summary[source] += 1
        if question.get("answer_check") and normalize_answer_check(question["answer_check"])[
            "needs_review"
        ]:
            summary["review_needed"] += 1
    return summary


def print_run_summary(summary: dict[str, int], output_dir: Path, title: str) -> None:
    print("处理汇总：", flush=True)
    print(f"- 总题数：{summary['total']}", flush=True)
    print(f"- 使用缓存：{summary['cache']}", flush=True)
    print(f"- 新生成解析：{summary['ai']}", flush=True)
    print(f"- 处理失败：{summary['failed']}", flush=True)
    print(f"- 需复核：{summary['review_needed']}", flush=True)
    print(f"- Word：{output_dir / f'{title}.docx'}", flush=True)
    print(f"- Markdown：{output_dir / f'{title}.md'}", flush=True)
    print(f"- 复核清单：{output_dir / 'review-needed.md'}", flush=True)


def render_markdown(questions: list[dict], title: str) -> str:
    lines = [f"# {title}", ""]
    current_course = None
    for index, question in enumerate(questions, 1):
        course = question.get("courseName") or "未命名课程"
        if course != current_course:
            lines.extend([f"## {course}", ""])
            current_course = course
        lines.extend(
            [
                f"### {index}. {question.get('question', '')}",
                "",
                f"> **答案：{question.get('answer', '')}**",
                f"> 题型：{question.get('type', '')}",
                "",
            ]
        )
        for option in question.get("options", []):
            lines.append(f"- {format_option_markdown(option, question.get('answer', ''))}")
        if question.get("options"):
            lines.append("")
        source = question.get("explanation_source", "unknown")
        explanation_lines = render_explanation_markdown(question.get("explanation"))
        answer_check_lines = render_answer_check_markdown(question.get("answer_check"))
        lines.extend(
            [
                "**解析**",
                "",
                *answer_check_lines,
                *explanation_lines,
                f"解析来源：{_source_label(source)}",
                "",
                "---",
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def render_answer_check_markdown(answer_check: object) -> list[str]:
    if not answer_check:
        return []
    item = normalize_answer_check(answer_check)
    if not item["needs_review"] and item["risk_level"] == "low":
        return []
    return [
        "**答案可能需要复核**",
        "这道题的导出答案与模型校验结果存在差异或不确定，详见 review-needed.md。",
        "",
    ]


def render_explanation_markdown(explanation: object) -> list[str]:
    item = normalize_explanation(explanation)
    lines = [
        "**为什么选：**",
        item["correct_reason"] or "暂无解析。",
        "",
    ]
    if item["wrong_options"]:
        lines.extend(["**为什么不选：**", ""])
        for wrong in item["wrong_options"]:
            prefix = f"{wrong['option']}：" if wrong["option"] else ""
            lines.append(f"- {prefix}{wrong['reason']}")
        lines.append("")
    if item["review_tip"]:
        lines.extend(["**复习抓手：**", f"> {item['review_tip']}", ""])
    if item["knowledge_points"]:
        lines.extend(["**知识补充：**", ""])
        lines.extend([f"- {point}" for point in item["knowledge_points"]])
        lines.append("")
    if item["principles"]:
        lines.extend(["**同类题判断法：**", ""])
        lines.extend([f"- {principle}" for principle in item["principles"]])
        lines.append("")
    return lines


def render_review_needed_markdown(questions: list[dict], title: str) -> str:
    flagged = [
        question
        for question in questions
        if question.get("answer_check")
        and normalize_answer_check(question["answer_check"])["needs_review"]
    ]
    lines = [f"# {title}", "", f"需复核题目数：{len(flagged)}", ""]
    for index, question in enumerate(flagged, 1):
        check = normalize_answer_check(question["answer_check"])
        lines.extend(
            [
                f"## {index}. {question.get('question', '')}",
                "",
                f"课程：{question.get('courseName', '未命名课程')}",
                "",
                f"导出答案：{check['provided_answer'] or question.get('answer', '')}",
                "",
                f"模型判断：{check['model_answer'] or '未提供'}",
                "",
                f"风险等级：{check['risk_level']}，置信度：{check['confidence']:.2f}",
                "",
                f"理由：{check['reason'] or '未提供'}",
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def format_option_markdown(option: str, answer: str) -> str:
    if is_correct_option(option, answer):
        return f"**{option}** ✅"
    return option


def is_correct_option(option: str, answer: str) -> bool:
    label_match = re.match(r"^([A-Z])\.\s*(.*)$", option.strip())
    option_label = label_match.group(1) if label_match else ""
    option_text = re.sub(r"^[A-Z]\.\s*", "", option).strip()
    answer_parts = [
        re.sub(r"^[A-Z]\.\s*", "", part.strip())
        for part in re.split(r"[；;,，、\s]+", str(answer))
        if part.strip()
    ]
    return any(part == option_label or part == option_text for part in answer_parts)


def _source_label(source: str) -> str:
    return {"ai": "AI", "platform": "平台", "cache": "缓存", "missing": "未生成"}.get(
        source, source
    )


def write_docx(questions: list[dict], title: str, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required to write DOCX files.") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document_styles(document)
    document.add_heading(title, level=0)
    current_course = None
    for index, question in enumerate(questions, 1):
        course = question.get("courseName") or "未命名课程"
        if course != current_course:
            document.add_heading(course, level=1)
            current_course = course
        question_paragraph = document.add_paragraph()
        question_run = question_paragraph.add_run(f"{index}. {question.get('question', '')}")
        question_run.bold = True
        question_run.font.size = Pt(QUESTION_FONT_SIZE_PT)
        question_run.font.color.rgb = RGBColor(46, 91, 170)
        answer_paragraph = document.add_paragraph()
        answer_run = answer_paragraph.add_run(f"答案：{question.get('answer', '')}")
        answer_run.bold = True
        answer_run.font.size = Pt(CALLOUT_FONT_SIZE_PT)
        answer_run.font.color.rgb = RGBColor(0, 0, 0)
        answer_paragraph.add_run(f"\n题型：{question.get('type', '')}")
        for option in question.get("options", []):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(option)
            if is_correct_option(option, question.get("answer", "")):
                run.bold = True
                run.font.color.rgb = RGBColor(34, 139, 34)
                paragraph.add_run("  ✅")
        add_answer_check_docx(document, question.get("answer_check"))
        add_explanation_docx(document, question.get("explanation"))
        document.add_paragraph(f"解析来源：{_source_label(question.get('explanation_source', 'unknown'))}")
        add_separator(document)
    document.save(output_path)


def configure_document_styles(document) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    font_name = docx_font_name()
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def add_answer_check_docx(document, answer_check: object) -> None:
    if not answer_check:
        return
    item = normalize_answer_check(answer_check)
    if not item["needs_review"] and item["risk_level"] == "low":
        return
    add_callout_paragraph(document, "答案可能需要复核")
    add_body_paragraph(document, "这道题的导出答案与模型校验结果存在差异或不确定，详见 review-needed.md。")


def add_explanation_docx(document, explanation: object) -> None:
    from docx.shared import Pt, RGBColor

    item = normalize_explanation(explanation)
    add_callout_paragraph(document, "解析")
    add_callout_paragraph(document, "为什么选：")
    add_body_paragraph(document, item["correct_reason"] or "暂无解析。")
    if item["wrong_options"]:
        add_callout_paragraph(document, "为什么不选：")
        for wrong in item["wrong_options"]:
            prefix = f"{wrong['option']}：" if wrong["option"] else ""
            add_list_paragraph(document, f"{prefix}{wrong['reason']}")
    if item["review_tip"]:
        add_callout_paragraph(document, "复习抓手：")
        tip = add_body_paragraph(document, item["review_tip"])
        shade_paragraph(tip, TIP_SHADE)
    if item["knowledge_points"]:
        add_callout_paragraph(document, "知识补充：")
        for point in item["knowledge_points"]:
            add_list_paragraph(document, point)
    if item["principles"]:
        add_callout_paragraph(document, "同类题判断法：")
        for principle in item["principles"]:
            add_list_paragraph(document, principle)


def add_callout_paragraph(document, text: str) -> None:
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(CALLOUT_FONT_SIZE_PT)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body_paragraph(document, text: str):
    from docx.shared import Pt

    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.18
    return paragraph


def add_list_paragraph(document, text: str):
    from docx.shared import Pt

    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def shade_paragraph(paragraph, fill: str) -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    paragraph._p.get_or_add_pPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{escape(fill)}"/>')
    )


def add_separator(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), LINE_COLOR)
    borders.append(bottom)
    p_pr.append(borders)


def build_outputs(args: argparse.Namespace) -> list[dict]:
    input_path = Path(args.input)
    output_dir = Path(args.output_dir)
    title = args.title or input_path.name or "学习通作业复习资料"
    cache_path = Path(args.cache) if args.cache else output_dir / "explanations.cache.json"
    load_dotenv(Path.cwd() / ".env")
    if input_path.is_dir():
        load_dotenv(input_path / ".env")

    questions = load_questions(input_path)
    questions = apply_limit(questions, args.limit)
    cache = load_cache(cache_path)
    cached_count = sum(1 for question in questions if question_key(question) in cache)
    print(f"输入题目：{len(questions)}", flush=True)
    print(f"已有缓存：{cached_count}", flush=True)
    print(f"开启答案校验：{'是' if args.verify_answers and not args.dry_run else '否'}", flush=True)

    def write_progress_cache(progress_cache: dict[str, dict], processed: list[dict]) -> None:
        save_json(cache_path, progress_cache)
        save_json(output_dir / "questions.partial.json", processed)

    enriched = enrich_questions(
        questions,
        cache=cache,
        dry_run=args.dry_run,
        verify_answers=args.verify_answers and not args.dry_run,
        cache_writer=write_progress_cache,
    )

    output_dir.mkdir(parents=True, exist_ok=True)
    save_json(output_dir / "questions.enriched.json", enriched)
    save_json(cache_path, update_cache(cache, enriched))
    (output_dir / "review-needed.md").write_text(
        render_review_needed_markdown(enriched, f"{title}-需复核题目"),
        encoding="utf-8",
    )
    (output_dir / f"{title}.md").write_text(
        render_markdown(enriched, title), encoding="utf-8"
    )
    write_docx(enriched, title, output_dir / f"{title}.docx")
    print_run_summary(build_run_summary(enriched), output_dir, title)
    return enriched


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Merge Chaoxing JSON homework exports and build review DOCX/Markdown."
    )
    parser.add_argument("input", help="JSON file or directory containing JSON exports.")
    parser.add_argument("--output-dir", default="output", help="Output directory.")
    parser.add_argument("--cache", help="Explanation cache JSON path.")
    parser.add_argument("--title", help="Document title and output filename stem.")
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Do not call the AI API; mark missing explanations instead.",
    )
    parser.add_argument(
        "--verify-answers",
        action="store_true",
        help="Ask the model to independently check exported answers and flag risks.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        help="Process only the first N questions. Useful for testing API output.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    enriched = build_outputs(args)
    print(f"Processed {len(enriched)} questions.")


if __name__ == "__main__":
    main()
