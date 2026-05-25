import json
import os
import tempfile
import unittest
from argparse import Namespace
from pathlib import Path

from scripts import homework_review


class HomeworkReviewTests(unittest.TestCase):
    def test_loads_chaoxing_json_files_and_deduplicates_questions(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            first = {
                "meta": {"courseName": "人工智能理论"},
                "questions": [
                    {
                        "courseName": "人工智能理论",
                        "type": "单选题",
                        "question": "1. 人工智能的英文缩写是？",
                        "options": ["A. AI ", " B. BI"],
                        "answer": "AI",
                        "score": 5,
                    }
                ],
            }
            second = {
                "meta": {"courseName": "人工智能理论"},
                "questions": [
                    {
                        "courseName": "人工智能理论",
                        "type": "单选题",
                        "question": "人工智能的英文缩写是？",
                        "options": ["A. AI", "B. BI"],
                        "answer": "AI",
                        "score": 5,
                    },
                    {
                        "courseName": "人工智能理论",
                        "type": "判断题",
                        "question": "机器学习是人工智能的重要分支。",
                        "options": [],
                        "answer": "正确",
                    },
                ],
            }
            (root / "first.json").write_text(
                json.dumps(first, ensure_ascii=False), encoding="utf-8"
            )
            (root / "second.json").write_text(
                json.dumps(second, ensure_ascii=False), encoding="utf-8"
            )

            questions = homework_review.load_questions(root)

            self.assertEqual(len(questions), 2)
            self.assertEqual(questions[0]["question"], "人工智能的英文缩写是？")
            self.assertEqual(questions[0]["options"], ["A. AI", "B. BI"])

    def test_load_questions_skips_generated_output_json_files(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            export = {
                "meta": {"courseName": "课程"},
                "questions": [{"type": "判断题", "question": "原始题", "answer": "正确"}],
            }
            generated = [{"question": "生成题", "answer": "错误"}]
            (root / "export.json").write_text(
                json.dumps(export, ensure_ascii=False), encoding="utf-8"
            )
            output = root / "output"
            output.mkdir()
            (output / "questions.enriched.json").write_text(
                json.dumps(generated, ensure_ascii=False), encoding="utf-8"
            )
            (output / "questions.partial.json").write_text(
                json.dumps(generated, ensure_ascii=False), encoding="utf-8"
            )

            questions = homework_review.load_questions(root)

            self.assertEqual(len(questions), 1)
            self.assertEqual(questions[0]["question"], "原始题")

    def test_load_questions_reads_collected_raw_json_under_output(self):
        with tempfile.TemporaryDirectory() as tmp:
            raw = Path(tmp) / "output" / "人工智能基础" / "raw"
            raw.mkdir(parents=True)
            (raw / "混合智能.json").write_text(
                json.dumps(
                    {
                        "meta": {"courseName": "人工智能基础"},
                        "questions": [
                            {"type": "单选题", "question": "机器人最基本的定义是？", "answer": "C"}
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            questions = homework_review.load_questions(raw)

            self.assertEqual(len(questions), 1)
            self.assertEqual(questions[0]["courseName"], "人工智能基础")

    def test_normalize_question_text_strips_inline_options_and_answers(self):
        text = (
            "机器人最基本的定义是（ ） A. 一种机械设备 B. 一种只能执行固定程序的机器 "
            "C. 一种能够自主或半自主执行任务的系统 D. 一种人工智能算法 "
            "C :一种能够自主或半自主执行任务的系统; 10 分"
        )

        self.assertEqual(
            homework_review.normalize_question_text(text),
            "机器人最基本的定义是（ ）",
        )

    def test_normalize_question_text_strips_trailing_option_label(self):
        self.assertEqual(
            homework_review.normalize_question_text("下列关于深度学习的描述，正确的是： A."),
            "下列关于深度学习的描述，正确的是：",
        )

    def test_load_questions_sorts_files_by_chapter_number(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            for name in ["第七章作业.json", "第一章作业.json", "第三章 总线作业.json"]:
                data = {
                    "meta": {"courseName": Path(name).stem},
                    "questions": [{"type": "判断题", "question": Path(name).stem, "answer": "正确"}],
                }
                (root / name).write_text(
                    json.dumps(data, ensure_ascii=False), encoding="utf-8"
                )

            questions = homework_review.load_questions(root)

            self.assertEqual(
                [question["question"] for question in questions],
                ["第一章作业", "第三章 总线作业", "第七章作业"],
            )

    def test_extracts_chapter_number_from_chinese_or_arabic_numerals(self):
        self.assertEqual(homework_review.extract_chapter_number("第十二章作业.json"), 12)
        self.assertEqual(homework_review.extract_chapter_number("第20章作业.json"), 20)
        self.assertIsNone(homework_review.extract_chapter_number("总线作业.json"))

    def test_generates_markdown_with_ai_explanation_marker(self):
        question = {
            "courseName": "人工智能理论",
            "type": "单选题",
            "question": "人工智能的英文缩写是？",
            "options": ["A. AI", "B. BI"],
            "answer": "AI",
            "explanation": {
                "correct_reason": "AI 是 Artificial Intelligence 的缩写。",
                "wrong_options": [
                    {"option": "B. BI", "reason": "BI 通常指商业智能，不是人工智能。"}
                ],
                "review_tip": "看到 Artificial Intelligence 就对应 AI。",
                "knowledge_points": ["人工智能英文缩写"],
                "principles": ["缩写题需要对应中英文概念。"],
            },
            "explanation_source": "ai",
        }

        markdown = homework_review.render_markdown([question], "复习资料")

        self.assertIn("# 复习资料", markdown)
        self.assertIn("题型：单选题", markdown)
        self.assertIn("**答案：AI**", markdown)
        self.assertIn("A. AI", markdown)
        self.assertIn("为什么选", markdown)
        self.assertIn("为什么不选", markdown)
        self.assertIn("复习抓手", markdown)
        self.assertIn("知识补充", markdown)
        self.assertIn("同类题判断法", markdown)
        self.assertIn("> **答案：AI**", markdown)
        self.assertIn("> 看到 Artificial Intelligence 就对应 AI。", markdown)
        self.assertIn("---", markdown)
        self.assertIn("- **A. AI** ✅", markdown)
        self.assertIn("解析来源：AI", markdown)
        self.assertLess(markdown.index("**答案：AI**"), markdown.index("题型：单选题"))
        self.assertLess(markdown.index("题型：单选题"), markdown.index("- **A. AI** ✅"))
        self.assertLess(markdown.index("- **A. AI** ✅"), markdown.index("**解析**"))

    def test_dry_run_adds_placeholder_without_api_call(self):
        questions = [
            {
                "courseName": "人工智能理论",
                "type": "判断题",
                "question": "机器学习是人工智能的重要分支。",
                "options": [],
                "answer": "正确",
            }
        ]

        enriched = homework_review.enrich_questions(
            questions,
            client=lambda _: (_ for _ in ()).throw(AssertionError("called API")),
            dry_run=True,
            logger=lambda _: None,
        )

        self.assertEqual(enriched[0]["explanation_source"], "missing")
        self.assertIn("待生成", enriched[0]["explanation"]["correct_reason"])

    def test_parses_structured_ai_explanation_json(self):
        raw = json.dumps(
            {
                "correct_reason": "非侵入式脑机接口通过头皮采集信号。",
                "wrong_options": [
                    {"option": "A", "reason": "侵入式需要植入电极。"}
                ],
                "review_tip": "看到头皮采集、无需手术，就优先判断为非侵入式。",
                "knowledge_points": ["脑机接口分类"],
                "principles": ["按是否植入电极区分侵入程度。"],
            },
            ensure_ascii=False,
        )

        explanation = homework_review.parse_explanation_response(raw)

        self.assertEqual(
            explanation["correct_reason"], "非侵入式脑机接口通过头皮采集信号。"
        )
        self.assertEqual(explanation["wrong_options"][0]["option"], "A")
        self.assertEqual(
            explanation["review_tip"], "看到头皮采集、无需手术，就优先判断为非侵入式。"
        )
        self.assertEqual(explanation["knowledge_points"], ["脑机接口分类"])

    def test_prompt_asks_for_student_facing_expanded_review_notes(self):
        messages = homework_review.build_prompt(
            {
                "type": "单选题",
                "question": "脑机接口按侵入程度如何分类？",
                "options": ["A. 侵入式", "B. 非侵入式"],
                "answer": "非侵入式",
            }
        )
        prompt_text = "\n".join(message["content"] for message in messages)

        self.assertIn("面向正在复习的学生", prompt_text)
        self.assertIn("不要只写标签", prompt_text)
        self.assertIn("展开说明", prompt_text)
        self.assertIn("review_tip", prompt_text)

    def test_correct_option_matching_does_not_match_substrings(self):
        self.assertFalse(
            homework_review.is_correct_option("A. 侵入式脑机接口", "非侵入式脑机接口")
        )
        self.assertTrue(
            homework_review.is_correct_option("B. 非侵入式脑机接口", "非侵入式脑机接口")
        )
        self.assertTrue(homework_review.is_correct_option("B. 非侵入式脑机接口", "B"))

    def test_parse_answer_check_json(self):
        raw = json.dumps(
            {
                "provided_answer": "A",
                "model_answer": "B",
                "verdict": "disagree",
                "confidence": 0.82,
                "risk_level": "high",
                "reason": "题干关键词更符合 B。",
                "needs_review": True,
            },
            ensure_ascii=False,
        )

        result = homework_review.parse_answer_check_response(raw)

        self.assertEqual(result["verdict"], "disagree")
        self.assertEqual(result["risk_level"], "high")
        self.assertTrue(result["needs_review"])

    def test_review_needed_markdown_lists_flagged_questions(self):
        questions = [
            {
                "courseName": "混合智能",
                "question": "测试题？",
                "type": "单选题",
                "options": ["A. 选项一", "B. 选项二"],
                "answer": "A",
                "answer_check": {
                    "provided_answer": "A",
                    "model_answer": "B",
                    "verdict": "disagree",
                    "confidence": 0.82,
                    "risk_level": "high",
                    "reason": "题干关键词更符合 B。",
                    "needs_review": True,
                },
            }
        ]

        markdown = homework_review.render_review_needed_markdown(questions, "复核清单")

        self.assertIn("# 复核清单", markdown)
        self.assertIn("测试题？", markdown)
        self.assertIn("题型：单选题", markdown)
        self.assertIn("- A. 选项一", markdown)
        self.assertIn("- B. 选项二", markdown)
        self.assertIn("导出答案：A", markdown)
        self.assertIn("模型判断：B", markdown)
        self.assertIn("风险等级：high", markdown)
        self.assertIn("判断状态：disagree", markdown)

    def test_low_risk_answer_check_is_hidden_from_review_material(self):
        lines = homework_review.render_answer_check_markdown(
            {
                "provided_answer": "A",
                "model_answer": "A",
                "verdict": "agree",
                "confidence": 0.95,
                "risk_level": "low",
                "reason": "一致。",
                "needs_review": False,
            }
        )

        self.assertEqual(lines, [])

    def test_high_risk_answer_check_is_short_warning_in_review_material(self):
        lines = homework_review.render_answer_check_markdown(
            {
                "provided_answer": "A",
                "model_answer": "B",
                "verdict": "disagree",
                "confidence": 0.82,
                "risk_level": "high",
                "reason": "题干关键词更符合 B。",
                "needs_review": True,
            }
        )

        rendered = "\n".join(lines)
        self.assertIn("答案可能需要复核", rendered)
        self.assertIn("详见 review-needed.md", rendered)
        self.assertNotIn("校验理由", rendered)
        self.assertNotIn("模型判断：B", rendered)

    def test_high_risk_answer_check_is_added_to_docx(self):
        from docx import Document

        document = Document()
        homework_review.add_answer_check_docx(
            document,
            {
                "provided_answer": "A",
                "model_answer": "B",
                "verdict": "disagree",
                "confidence": 0.82,
                "risk_level": "high",
                "reason": "题干关键词更符合 B。",
                "needs_review": True,
            },
        )

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("答案可能需要复核", text)
        self.assertIn("详见 review-needed.md", text)

    def test_answer_visibility_risk_is_shown_as_short_warning(self):
        lines = homework_review.render_answer_source_markdown(
            {"answer_visibility": "student_answer_only"}
        )

        rendered = "\n".join(lines)
        self.assertIn("答案来源需要留意", rendered)
        self.assertIn("详见 review-needed.md", rendered)

    def test_review_needed_markdown_includes_answer_visibility_risk(self):
        markdown = homework_review.render_review_needed_markdown(
            [
                {
                    "courseName": "汇编语言",
                    "question": "只显示我的答案的题？",
                    "type": "判断题",
                    "options": ["A. 对", "B. 错"],
                    "answer": "对",
                    "answer_visibility": "student_answer_only",
                    "student_answer": "对",
                    "correct_answer": "",
                }
            ],
            "复核清单",
        )

        self.assertIn("只显示我的答案的题？", markdown)
        self.assertIn("答案来源：student_answer_only", markdown)
        self.assertIn("导出答案：对", markdown)

    def test_docx_body_paragraph_helper_applies_indent_and_spacing(self):
        from docx import Document

        document = Document()
        paragraph = homework_review.add_body_paragraph(document, "正文")

        self.assertEqual(paragraph.paragraph_format.left_indent.pt, 12)
        self.assertEqual(paragraph.paragraph_format.space_after.pt, 8)

    def test_docx_font_defaults_to_microsoft_yahei(self):
        old_font = os.environ.pop("DOCX_FONT", None)
        try:
            self.assertEqual(homework_review.docx_font_name(), "Microsoft YaHei")
        finally:
            if old_font is not None:
                os.environ["DOCX_FONT"] = old_font

    def test_docx_font_uses_valid_environment_value(self):
        old_font = os.environ.get("DOCX_FONT")
        os.environ["DOCX_FONT"] = "Maple Mono"
        try:
            self.assertEqual(homework_review.docx_font_name(), "Maple Mono")
        finally:
            if old_font is None:
                os.environ.pop("DOCX_FONT", None)
            else:
                os.environ["DOCX_FONT"] = old_font

    def test_docx_font_falls_back_when_environment_value_is_invalid(self):
        old_font = os.environ.get("DOCX_FONT")
        os.environ["DOCX_FONT"] = "Bad/Font<Name>"
        try:
            self.assertEqual(homework_review.docx_font_name(), "Microsoft YaHei")
        finally:
            if old_font is None:
                os.environ.pop("DOCX_FONT", None)
            else:
                os.environ["DOCX_FONT"] = old_font

    def test_applies_limit_before_processing_outputs(self):
        questions = [{"question": str(index)} for index in range(5)]

        limited = homework_review.apply_limit(questions, 2)

        self.assertEqual([item["question"] for item in limited], ["0", "1"])

    def test_log_progress_writes_stage_and_index(self):
        messages = []

        homework_review.log_progress("生成解析", 2, 10, "这是一道很长很长的题目", logger=messages.append)

        self.assertIn("[2/10] 生成解析：这是一道很长很长的题目", messages[0])

    def test_enrich_questions_writes_cache_after_each_question(self):
        cache_snapshots = []
        questions = [
            {"type": "判断题", "question": "题目一", "answer": "正确", "options": []},
            {"type": "判断题", "question": "题目二", "answer": "错误", "options": []},
        ]

        homework_review.enrich_questions(
            questions,
            client=lambda _: json.dumps(
                {
                    "correct_reason": "理由",
                    "wrong_options": [],
                    "review_tip": "抓手",
                    "knowledge_points": [],
                    "principles": [],
                },
                ensure_ascii=False,
            ),
            cache={},
            cache_writer=lambda cache, processed: cache_snapshots.append(
                (len(cache), len(processed))
            ),
            logger=lambda _: None,
        )

        self.assertEqual(cache_snapshots, [(1, 1), (2, 2)])

    def test_save_json_atomic_writes_file_and_removes_temp_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "cache.json"

            homework_review.save_json_atomic(path, {"ok": True})

            self.assertEqual(json.loads(path.read_text(encoding="utf-8")), {"ok": True})
            self.assertFalse(path.with_suffix(path.suffix + ".tmp").exists())

    def test_enrich_questions_continues_when_one_question_fails(self):
        calls = {"count": 0}
        questions = [
            {"type": "判断题", "question": "题目一", "answer": "正确", "options": []},
            {"type": "判断题", "question": "题目二", "answer": "错误", "options": []},
        ]

        def flaky_client(_):
            calls["count"] += 1
            if calls["count"] == 1:
                raise RuntimeError("api failed")
            return json.dumps(
                {
                    "correct_reason": "理由",
                    "wrong_options": [],
                    "review_tip": "抓手",
                    "knowledge_points": [],
                    "principles": [],
                },
                ensure_ascii=False,
            )

        enriched = homework_review.enrich_questions(
            questions,
            client=flaky_client,
            cache={},
            logger=lambda _: None,
        )

        self.assertEqual(enriched[0]["explanation_source"], "failed")
        self.assertIn("api failed", enriched[0]["processing_error"])
        self.assertEqual(enriched[1]["explanation_source"], "ai")

    def test_build_run_summary_counts_cache_ai_failures_and_review_needed(self):
        summary = homework_review.build_run_summary(
            [
                {"explanation_source": "cache"},
                {"explanation_source": "ai", "answer_check": {"needs_review": True}},
                {"explanation_source": "failed"},
            ]
        )

        self.assertEqual(summary["total"], 3)
        self.assertEqual(summary["cache"], 1)
        self.assertEqual(summary["ai"], 1)
        self.assertEqual(summary["failed"], 1)
        self.assertEqual(summary["review_needed"], 1)

    def test_cached_explanation_is_marked_as_cache_source(self):
        question = {"type": "判断题", "question": "题目", "answer": "正确", "options": []}
        cache = {
            homework_review.question_key(question): {
                "explanation": {"correct_reason": "缓存解析"},
                "explanation_source": "ai",
            }
        }

        enriched = homework_review.enrich_questions(
            [question],
            cache=cache,
            logger=lambda _: None,
        )

        self.assertEqual(enriched[0]["explanation_source"], "cache")
        self.assertEqual(enriched[0]["cached_explanation_source"], "ai")

    def test_cleans_multiselect_answers_and_duplicate_option_labels(self):
        raw = {
            "type": "多选题",
            "question": "深度神经网络的训练难点包括：",
            "options": ["A. A. 梯度消失", "B. B. 参数多", " C. 训练慢"],
            "answer": "梯度消失###参数多###训练慢",
        }

        cleaned = homework_review.normalize_question(raw)

        self.assertEqual(cleaned["options"], ["A. 梯度消失", "B. 参数多", "C. 训练慢"])
        self.assertEqual(cleaned["answer"], "梯度消失；参数多；训练慢")

    def test_loads_dotenv_without_overriding_existing_environment(self):
        with tempfile.TemporaryDirectory() as tmp:
            env_file = Path(tmp) / ".env"
            env_file.write_text(
                'AI_API_KEY="from-file"\nAI_MODEL=deepseek-v4-flash\n',
                encoding="utf-8",
            )
            old_key = os.environ.get("AI_API_KEY")
            old_model = os.environ.get("AI_MODEL")
            os.environ["AI_API_KEY"] = "from-env"
            os.environ.pop("AI_MODEL", None)
            try:
                homework_review.load_dotenv(env_file)

                self.assertEqual(os.environ["AI_API_KEY"], "from-env")
                self.assertEqual(os.environ["AI_MODEL"], "deepseek-v4-flash")
            finally:
                if old_key is None:
                    os.environ.pop("AI_API_KEY", None)
                else:
                    os.environ["AI_API_KEY"] = old_key
                if old_model is None:
                    os.environ.pop("AI_MODEL", None)
                else:
                    os.environ["AI_MODEL"] = old_model


if __name__ == "__main__":
    unittest.main()
